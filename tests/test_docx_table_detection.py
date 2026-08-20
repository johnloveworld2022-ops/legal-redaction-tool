from core.convert import convert_docx


def test_docx_with_native_table_flagged_possible_table(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("以下为当事人信息:")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "身份证号"
    table.cell(1, 0).text = "张三"
    table.cell(1, 1).text = "110101199003072316"
    out_docx = tmp_path / "with_table.docx"
    doc.save(out_docx)

    text, _images, has_table = convert_docx(out_docx)
    assert has_table is True


def test_docx_without_table_not_flagged(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("普通合同正文,没有表格。")
    out_docx = tmp_path / "no_table.docx"
    doc.save(out_docx)

    text, _images, has_table = convert_docx(out_docx)
    assert has_table is False
