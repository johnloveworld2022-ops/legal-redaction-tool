import subprocess

import pytest
from PIL import Image, ImageDraw, ImageFont

from core.convert import convert_docx, convert_pdf
from core.subprocess_utils import SubprocessTimeoutError

CHINESE_FONT = "/System/Library/Fonts/STHeiti Light.ttc"


def test_one_failing_pdf_page_does_not_abort_the_rest(tmp_path, monkeypatch):
    # 2-page scanned PDF; force the render step to fail only on page 1
    img1 = Image.new("RGB", (900, 100), "white")
    d1 = ImageDraw.Draw(img1)
    font = ImageFont.truetype(CHINESE_FONT, 40)
    d1.text((20, 20), "第一页内容", font=font, fill="black")

    img2 = Image.new("RGB", (900, 100), "white")
    d2 = ImageDraw.Draw(img2)
    d2.text((20, 20), "第二页原告张三", font=font, fill="black")

    out_pdf = tmp_path / "two_page.pdf"
    img1.save(out_pdf, "PDF", save_all=True, append_images=[img2])

    import core.convert as convert_module

    real_render = convert_module._render_pdf_page_to_png

    def flaky_render(path, page_num, out_dir):
        if page_num == 1:
            raise SubprocessTimeoutError(["pdftoppm"], 120)
        return real_render(path, page_num, out_dir)

    monkeypatch.setattr(convert_module, "_render_pdf_page_to_png", flaky_render)

    pages = convert_pdf(out_pdf)
    assert len(pages) == 2
    assert pages[0].source == "failed"
    assert pages[0].failure_reason is not None
    assert pages[1].source == "ocr"
    assert "张三" in pages[1].text


def test_corrupt_docx_embedded_image_recorded_as_failed_not_silently_dropped(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("正文内容")
    out_docx = tmp_path / "case.docx"
    doc.save(out_docx)

    # manually inject a corrupt "image" into word/media/ by editing the
    # zip after python-docx writes it, since add_picture() would refuse a
    # non-image file outright
    import zipfile

    tmp_docx2 = tmp_path / "case_corrupt.docx"
    with zipfile.ZipFile(out_docx) as zin, zipfile.ZipFile(tmp_docx2, "w") as zout:
        for item in zin.infolist():
            zout.writestr(item, zin.read(item.filename))
        zout.writestr("word/media/image1.png", b"not actually a png")

    text, image_pages, _has_table = convert_docx(tmp_docx2)
    assert len(image_pages) == 1
    assert image_pages[0].source == "failed"
    assert image_pages[0].failure_reason is not None
