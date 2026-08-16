import random
import subprocess

import pytest
from PIL import Image, ImageDraw, ImageFilter, ImageFont

from core.convert import convert_docx, convert_pdf, convert_to_text

CHINESE_FONT = "/System/Library/Fonts/STHeiti Light.ttc"
SENTENCE = "原告张三身份证号110101199003072316"


@pytest.fixture
def docx_file(tmp_path):
    src_txt = tmp_path / "src.txt"
    src_txt.write_text(SENTENCE, encoding="utf-8")
    out_docx = tmp_path / "case.docx"
    subprocess.run(
        ["textutil", "-convert", "docx", "-output", str(out_docx), str(src_txt)],
        check=True, capture_output=True,
    )
    return out_docx


@pytest.fixture
def docx_with_embedded_photo(tmp_path):
    # Simulates the real-world workflow: a photographed page inserted
    # directly into a Word document, with some typed text around it.
    # textutil's plain-text conversion only reads the document's text
    # runs -- it never looks at word/media/*, so before the docx-OCR fix
    # this photo's content was silently invisible to the whole pipeline,
    # not flagged, not OCR'd, nothing.
    from docx import Document

    img_path = tmp_path / "exhibit.png"
    img = Image.new("RGB", (900, 150), "white")
    d = ImageDraw.Draw(img)
    font = ImageFont.truetype(CHINESE_FONT, 40)
    d.text((20, 40), SENTENCE, font=font, fill="black")
    img.save(img_path)

    doc = Document()
    doc.add_paragraph("以下为本案证据材料:")
    doc.add_picture(str(img_path))
    out_docx = tmp_path / "case_with_photo.docx"
    doc.save(out_docx)
    return out_docx


@pytest.fixture
def text_layer_pdf(tmp_path):
    # cupsfilter renders a plain-text file to a real, selectable-text PDF
    # via the system print filter -- a proper text-layer PDF, not an image.
    src_txt = tmp_path / "src.txt"
    src_txt.write_text(SENTENCE, encoding="utf-8")
    out_pdf = tmp_path / "case.pdf"
    result = subprocess.run(
        ["cupsfilter", str(src_txt)], capture_output=True, check=True
    )
    out_pdf.write_bytes(result.stdout)
    return out_pdf


@pytest.fixture
def scanned_image_pdf(tmp_path):
    img = Image.new("RGB", (900, 150), "white")
    d = ImageDraw.Draw(img)
    font = ImageFont.truetype(CHINESE_FONT, 40)
    d.text((20, 40), SENTENCE, font=font, fill="black")
    out_pdf = tmp_path / "scanned.pdf"
    img.save(out_pdf, "PDF")
    return out_pdf


@pytest.fixture
def degraded_scanned_pdf(tmp_path):
    random.seed(42)
    img = Image.new("RGB", (900, 150), "white")
    d = ImageDraw.Draw(img)
    font = ImageFont.truetype(CHINESE_FONT, 40)
    d.text((20, 40), SENTENCE, font=font, fill=(135, 135, 135))
    img = img.filter(ImageFilter.GaussianBlur(2.5))
    px = img.load()
    w, h = img.size
    for x in range(w):
        for y in range(h):
            r, g, b = px[x, y]
            n = random.randint(-35, 35)
            px[x, y] = (
                max(0, min(255, r + n)), max(0, min(255, g + n)), max(0, min(255, b + n))
            )
    out_pdf = tmp_path / "degraded.pdf"
    img.save(out_pdf, "PDF")
    return out_pdf


@pytest.fixture
def composite_page_pdf(tmp_path):
    # Simulates the real-world case: a Word/PDF page with a short typed
    # caption (a genuine, short native text layer) AND an inserted
    # photographed exhibit on the SAME page. Built by overlaying an
    # image-only page onto a real cupsfilter-rendered text page via pypdf,
    # so the result has both a real extractable text layer and a real
    # embedded image object -- verified empirically (see session notes)
    # that pdftotext extracts the caption and pdfimages -list detects the
    # image with this specific page ordering.
    from pypdf import PdfReader, PdfWriter

    img = Image.new("RGB", (900, 150), "white")
    d = ImageDraw.Draw(img)
    font = ImageFont.truetype(CHINESE_FONT, 40)
    d.text((20, 40), SENTENCE, font=font, fill="black")
    img_pdf_path = tmp_path / "exhibit.pdf"
    img.save(img_pdf_path, "PDF")

    caption_txt = tmp_path / "caption.txt"
    # Must be >= MIN_TEXT_LAYER_CHARS so this page actually exercises the
    # vulnerable "trust the text layer, skip OCR" branch under the old
    # logic -- a too-short caption would fall into the OCR branch anyway
    # by coincidence and not reproduce the bug.
    caption_text = "本案原告提交的证据材料附卷如下,请予以审查核实,谢谢配合处理。"
    assert len(caption_text) >= 20
    caption_txt.write_text(caption_text, encoding="utf-8")
    caption_pdf_path = tmp_path / "caption.pdf"
    result = subprocess.run(["cupsfilter", str(caption_txt)], capture_output=True, check=True)
    caption_pdf_path.write_bytes(result.stdout)

    img_reader = PdfReader(str(img_pdf_path))
    caption_reader = PdfReader(str(caption_pdf_path))
    writer = PdfWriter()
    base_page = caption_reader.pages[0]
    base_page.merge_page(img_reader.pages[0])
    writer.add_page(base_page)

    out_pdf = tmp_path / "composite.pdf"
    with open(out_pdf, "wb") as f:
        writer.write(f)
    return out_pdf, caption_text


def test_docx_extracts_expected_text(docx_file):
    text, images = convert_docx(docx_file)
    assert SENTENCE in text
    assert images == []


def test_docx_with_embedded_photo_ocrs_it_instead_of_silently_dropping_it(
    docx_with_embedded_photo,
):
    text, images = convert_docx(docx_with_embedded_photo)
    assert "以下为本案证据材料" in text
    assert "身份证号" in text  # from OCR of the embedded photo
    assert len(images) == 1
    assert images[0].source == "ocr"


def test_docx_with_no_images_behaves_as_before(docx_file):
    text, images = convert_docx(docx_file)
    assert SENTENCE in text
    assert "OCR" not in text
    assert images == []


def test_docx_embedded_photo_image_persisted_when_output_dir_given(
    tmp_path, docx_with_embedded_photo
):
    image_dir = tmp_path / "saved"
    image_dir.mkdir()
    text, images = convert_docx(docx_with_embedded_photo, image_output_dir=image_dir)
    assert images[0].source_image_path is not None
    assert images[0].source_image_path.exists()
    assert images[0].source_image_path.parent == image_dir


def test_text_layer_pdf_uses_text_layer_not_ocr(text_layer_pdf):
    pages = convert_pdf(text_layer_pdf)
    assert len(pages) == 1
    assert pages[0].source == "text_layer"
    assert SENTENCE in pages[0].text


def test_scanned_image_pdf_falls_back_to_ocr(scanned_image_pdf):
    pages = convert_pdf(scanned_image_pdf)
    assert len(pages) == 1
    assert pages[0].source == "ocr"
    assert "张三" in pages[0].text
    assert "110101199003072316" in pages[0].text


def test_convert_to_text_marks_ocr_pages(scanned_image_pdf):
    text = convert_to_text(scanned_image_pdf)
    assert "(OCR)" in text
    assert "张三" in text


def test_unsupported_extension_raises(tmp_path):
    bad_file = tmp_path / "case.rtf"
    bad_file.write_text("hello", encoding="utf-8")
    with pytest.raises(ValueError):
        convert_to_text(bad_file)


def test_clean_ocr_page_has_no_low_confidence_lines(scanned_image_pdf):
    pages = convert_pdf(scanned_image_pdf)
    assert pages[0].low_confidence_lines == []


def test_degraded_ocr_page_is_flagged_for_manual_check(degraded_scanned_pdf):
    pages = convert_pdf(degraded_scanned_pdf)
    assert pages[0].source == "ocr"
    assert len(pages[0].low_confidence_lines) > 0


def test_ocr_page_image_persisted_when_output_dir_given(tmp_path, scanned_image_pdf):
    image_dir = tmp_path / "saved_pages"
    image_dir.mkdir()
    pages = convert_pdf(scanned_image_pdf, image_output_dir=image_dir)
    assert pages[0].source_image_path is not None
    assert pages[0].source_image_path.exists()
    assert pages[0].source_image_path.parent == image_dir


def test_text_layer_page_has_no_source_image(text_layer_pdf):
    pages = convert_pdf(text_layer_pdf)
    assert pages[0].source_image_path is None
    assert pages[0].low_confidence_lines == []


def test_composite_page_ocrs_image_instead_of_silently_skipping_it(composite_page_pdf):
    # The critical bug this guards against: a page with SOME native text
    # (here, a short caption) must not let that text layer stand in for
    # the whole page and skip OCR -- the embedded photographed exhibit on
    # the same page has to be OCR'd too, or its content silently vanishes
    # with no error and no low-confidence flag to catch it.
    pdf_path, caption_text = composite_page_pdf
    pages = convert_pdf(pdf_path)
    assert len(pages) == 1
    assert pages[0].source == "mixed"
    assert caption_text in pages[0].text
    # the embedded exhibit image must have actually been OCR'd, not
    # dropped -- some recognizable fragment of its content must appear
    assert "身份证号" in pages[0].text


def test_composite_page_saves_its_image_for_review(tmp_path, composite_page_pdf):
    pdf_path, _caption_text = composite_page_pdf
    image_dir = tmp_path / "saved"
    image_dir.mkdir()
    pages = convert_pdf(pdf_path, image_output_dir=image_dir)
    assert pages[0].source_image_path is not None
    assert pages[0].source_image_path.exists()


def test_pure_text_page_is_unaffected_by_the_image_check(text_layer_pdf):
    # Regression guard: a page with real text and NO embedded image must
    # still take the fast text-only path (no OCR run, no duplication).
    pages = convert_pdf(text_layer_pdf)
    assert pages[0].source == "text_layer"
