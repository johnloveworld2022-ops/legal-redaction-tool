import subprocess

import pytest
from PIL import Image, ImageDraw, ImageFont

import core.case_workspace as case_workspace
from core.mapping_store import MappingStore
from core.orchestrator import approve_case_export, process_case_files

CHINESE_FONT = "/System/Library/Fonts/STHeiti Light.ttc"


class FakeClient:
    def __init__(self, response='{"entities": []}'):
        self._response = response

    def generate(self, prompt: str) -> str:
        return self._response


@pytest.fixture(autouse=True)
def isolated_workbench(tmp_path, monkeypatch):
    monkeypatch.setattr(case_workspace, "WORKBENCH_ROOT", tmp_path / "工作台")
    yield


@pytest.fixture
def clean_txt_file(tmp_path):
    p = tmp_path / "clean.txt"
    p.write_text("本合同双方已就交货期达成一致,无需另行通知。", encoding="utf-8")
    return p


@pytest.fixture
def pii_txt_file(tmp_path):
    p = tmp_path / "pii.txt"
    p.write_text("原告张三身份证号110101199003072316,住址不详。", encoding="utf-8")
    return p


def _cleanup_keychain(case_name: str):
    from core.case_workspace import case_workspace_for
    ws = case_workspace_for(case_name)
    MappingStore(path=ws.mapping_path, keychain_service=ws.keychain_service).delete_key()


def test_clean_file_auto_exports(clean_txt_file):
    case_name = "测试案A"
    try:
        summary = process_case_files(case_name, [clean_txt_file], llm_client=FakeClient())
        assert summary.all_clean is True
        assert len(summary.exported_files) == 1
        assert summary.exported_files[0].exists()
    finally:
        _cleanup_keychain(case_name)


def test_regex_only_high_confidence_finding_auto_exports(pii_txt_file):
    # Only the ID card matches (regex, high confidence); the llm stage ran
    # and found nothing in this fake response, so nothing blocks export.
    case_name = "测试案B"
    try:
        summary = process_case_files(case_name, [pii_txt_file], llm_client=FakeClient())
        assert summary.all_clean is True  # regex-only high confidence, llm found nothing
        assert "⟦身份证001⟧" in summary.exported_files[0].read_text(encoding="utf-8")
    finally:
        _cleanup_keychain(case_name)


def test_llm_finding_blocks_and_report_written(pii_txt_file):
    case_name = "测试案C"
    llm_json = '{"entities": [{"text": "张三", "type": "PERSON"}]}'
    try:
        summary = process_case_files(case_name, [pii_txt_file], llm_client=FakeClient(llm_json))
        assert summary.all_clean is False
        assert summary.exported_files == []
        assert summary.report_path.exists()
        assert "张三" not in summary.report_path.read_text(encoding="utf-8")
    finally:
        _cleanup_keychain(case_name)


def test_approve_after_blocked_run_exports_files(pii_txt_file):
    case_name = "测试案D"
    llm_json = '{"entities": [{"text": "张三", "type": "PERSON"}]}'
    try:
        summary = process_case_files(case_name, [pii_txt_file], llm_client=FakeClient(llm_json))
        assert summary.all_clean is False

        approve_summary = approve_case_export(case_name)
        assert approve_summary.exported_count == 1
        assert len(list(approve_summary.approved_dir.iterdir())) == 1
    finally:
        _cleanup_keychain(case_name)


def test_approve_with_nothing_pending_reports_zero(tmp_path):
    case_name = "测试案E"
    from core.case_workspace import case_workspace_for
    case_workspace_for(case_name)  # create empty workspace, nothing processed
    try:
        approve_summary = approve_case_export(case_name)
        assert approve_summary.exported_count == 0
    finally:
        _cleanup_keychain(case_name)


def test_duplicate_lexicon_name_blocks_auto_export(clean_txt_file):
    case_name = "测试案G"
    from core.case_workspace import case_workspace_for

    ws = case_workspace_for(case_name)
    ws.lexicon_path.write_text("张三\n张三\n", encoding="utf-8")
    try:
        summary = process_case_files(case_name, [clean_txt_file], llm_client=FakeClient())
        assert summary.all_clean is False
        report_text = summary.report_path.read_text(encoding="utf-8")
        assert "张三" in report_text
        assert "同名" in report_text
    finally:
        _cleanup_keychain(case_name)


def test_docx_native_table_blocks_auto_export(tmp_path):
    from docx import Document

    case_name = "测试案H"
    doc = Document()
    doc.add_paragraph("以下为当事人信息:")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "身份证号"
    table.cell(1, 0).text = "张三"
    table.cell(1, 1).text = "110101199003072316"
    docx_path = tmp_path / "with_table.docx"
    doc.save(docx_path)
    try:
        summary = process_case_files(case_name, [docx_path], llm_client=FakeClient())
        assert summary.all_clean is False
        report_text = summary.report_path.read_text(encoding="utf-8")
        assert "表格" in report_text
    finally:
        _cleanup_keychain(case_name)


def test_page_processing_failure_blocks_auto_export(tmp_path, monkeypatch):
    img1 = Image.new("RGB", (900, 100), "white")
    d1 = ImageDraw.Draw(img1)
    font = ImageFont.truetype(CHINESE_FONT, 40)
    d1.text((20, 20), "第一页干净内容", font=font, fill="black")
    pdf_path = tmp_path / "one_page.pdf"
    img1.save(pdf_path, "PDF")

    import core.convert as convert_module

    def always_fail(path, page_num, out_dir):
        raise RuntimeError("模拟渲染失败")

    monkeypatch.setattr(convert_module, "_render_pdf_page_to_png", always_fail)

    case_name = "测试案I"
    try:
        summary = process_case_files(case_name, [pdf_path], llm_client=FakeClient())
        assert summary.all_clean is False
        report_text = summary.report_path.read_text(encoding="utf-8")
        assert "处理失败" in report_text
    finally:
        _cleanup_keychain(case_name)


def test_second_document_reuses_mapping_token_across_case(tmp_path):
    case_name = "测试案F"
    from core.case_workspace import case_workspace_for

    ws = case_workspace_for(case_name)
    ws.lexicon_path.write_text("张三\n", encoding="utf-8")

    doc1 = tmp_path / "doc1.txt"
    doc1.write_text("原告张三身份证号110101199003072316。", encoding="utf-8")
    doc2 = tmp_path / "doc2.txt"
    doc2.write_text("被告未回应张三的请求。", encoding="utf-8")
    try:
        s1 = process_case_files(case_name, [doc1], llm_client=FakeClient())
        assert s1.all_clean is True

        store = MappingStore(path=ws.mapping_path, keychain_service=ws.keychain_service)
        token = next(tok for tok, real in store.load().items() if real == "张三")

        s2 = process_case_files(case_name, [doc2], llm_client=FakeClient())
        assert s2.all_clean is True
        text2 = s2.exported_files[0].read_text(encoding="utf-8")
        assert token in text2
    finally:
        _cleanup_keychain(case_name)
